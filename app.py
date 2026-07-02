import os
import csv
import json
import uuid
from flask import Flask, request, jsonify, render_template, send_file, Response
from werkzeug.utils import secure_filename
import cv2
import numpy as np
import fitz

# Thêm path để import omr_processor
from omr_processor import OMRProcessor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__, template_folder="templates")

# Cấu hình tải lên
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Khởi tạo OMR Processor
processor = OMRProcessor(os.path.join(BASE_DIR, "data", "bubble_coordinates.json"))

# Danh sách sinh viên và Đáp án gốc nạp từ file khi chạy server
CANDIDATES_DB = {}   # CCCD -> { Tên, Năm sinh, Ngày cấp, Nơi thường trú }
ANSWER_KEY = {}      # Câu -> Đáp án chuẩn
PROCESSED_SHEETS = {} # sheet_id -> thông tin kết quả chấm thi

def load_candidates():
    global CANDIDATES_DB
    CANDIDATES_DB = {}
    excel_path = os.path.join(BASE_DIR, "data", "candidates.xlsx")
    
    if not os.path.exists(excel_path):
        # Tự động tạo tệp candidates.xlsx mẫu nếu chưa có
        os.makedirs(os.path.dirname(excel_path), exist_ok=True)
        try:
            import pandas as pd
            df_template = pd.DataFrame([
                ["Nguyễn Văn A", "1995", "012345678901", "01/01/2015", "Hà Nội"],
                ["Trần Thị B", "1998", "987654321098", "02/02/2018", "Hồ Chí Minh"],
                ["Lê Văn C", "2000", "111111111111", "03/03/2020", "Đà Nẵng"]
            ], columns=["Tên", "Năm sinh", "CCCD", "Ngày cấp", "Nơi thường trú"])
            
            # Ghi ra file Excel mẫu
            df_template.to_excel(excel_path, sheet_name="Lớp 1", index=False)
            print(f"[+] Đã tự động tạo tệp danh sách thí sinh Excel mẫu tại: {excel_path}")
        except Exception as e:
            print(f"[!] Không thể tự tạo tệp Excel mẫu: {e}")
            return
            
    try:
        import pandas as pd
        xls = pd.ExcelFile(excel_path)
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name)
            # Làm sạch headers
            df.columns = [str(c).strip() for c in df.columns]
            if "CCCD" in df.columns:
                for _, row in df.iterrows():
                    # Đọc CCCD dưới dạng chuỗi và làm sạch
                    cccd_raw = str(row["CCCD"]).strip()
                    if not cccd_raw or cccd_raw == "nan":
                        continue
                    # Bỏ phần thập phân .0 nếu Excel đọc nhầm thành float
                    cccd = cccd_raw.split(".")[0]
                    if cccd:
                        CANDIDATES_DB[cccd] = {
                            "name": str(row.get("Tên", "")).strip() if pd.notna(row.get("Tên")) else "",
                            "dob": str(row.get("Năm sinh", "")).strip().split(".")[0] if pd.notna(row.get("Năm sinh")) else "",
                            "date_issued": str(row.get("Ngày cấp", "")).strip() if pd.notna(row.get("Ngày cấp")) else "",
                            "address": str(row.get("Nơi thường trú", "")).strip() if pd.notna(row.get("Nơi thường trú")) else "",
                            "sheet_name": sheet_name
                        }
        print(f"[+] Loaded {len(CANDIDATES_DB)} candidates from Excel across all sheets.")
    except Exception as e:
        print(f"[!] Error loading candidates from Excel: {e}")

def load_answer_key():
    global ANSWER_KEY
    key_path = os.path.join(BASE_DIR, "data", "answer_key.json")
    if not os.path.exists(key_path):
        print(f"[!] Warning: File {key_path} không tồn tại.")
        return
    try:
        with open(key_path, "r", encoding="utf-8") as f:
            ANSWER_KEY = json.load(f)
        print(f"[+] Loaded answer key for {len(ANSWER_KEY)} questions.")
    except Exception as e:
        print(f"[!] Error loading answer key: {e}")

# Tải dữ liệu ban đầu
load_candidates()
load_answer_key()

def calculate_score(student_answers):
    """Tính số câu đúng và quy điểm hệ 10"""
    if not ANSWER_KEY:
        return 0, 0
    correct_count = 0
    for q_num, ans in ANSWER_KEY.items():
        student_ans = student_answers.get(str(q_num), "")
        if student_ans == ans:
            correct_count += 1
    # Tính điểm hệ 10 làm tròn 2 chữ số thập phân
    score = round((correct_count / len(ANSWER_KEY)) * 10, 2)
    return correct_count, score

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/candidates", methods=["GET"])
def get_candidates():
    # Trả về danh sách sinh viên dưới dạng mảng để hiển thị hỗ trợ gõ autocomplete/chọn sửa lỗi
    cand_list = [{"CCCD": k, **v} for k, v in CANDIDATES_DB.items()]
    return jsonify(cand_list)

@app.route("/api/answer-key", methods=["GET"])
def get_answer_key():
    return jsonify(ANSWER_KEY)

@app.route("/api/clear", methods=["POST"])
def clear_sessions():
    global PROCESSED_SHEETS
    PROCESSED_SHEETS = {}
    return jsonify({"status": "success", "message": "Đã làm sạch bộ nhớ tạm."})

@app.route("/api/upload", methods=["POST"])
def upload_files():
    global PROCESSED_SHEETS
    
    if 'files' not in request.files:
        return jsonify({"status": "error", "message": "Không tìm thấy file nào được tải lên."}), 400
        
    uploaded_files = request.files.getlist('files')
    results = []
    
    for file in uploaded_files:
        if not file.filename:
            continue
            
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Xác định là PDF hay hình ảnh
        _, ext = os.path.splitext(filename.lower())
        
        pages_to_process = []
        if ext == '.pdf':
            try:
                pages_to_process = processor.convert_pdf_to_images(file_path)
            except Exception as e:
                print(f"[!] Lỗi chuyển đổi PDF: {e}")
                results.append({
                    "filename": filename,
                    "error": f"Lỗi đọc file PDF: {str(e)}"
                })
                continue
        else:
            # Là ảnh JPG, PNG...
            try:
                img = cv2.imread(file_path)
                pages_to_process = [(1, img)]
            except Exception as e:
                results.append({
                    "filename": filename,
                    "error": f"Lỗi đọc file ảnh: {str(e)}"
                })
                continue
                
        # Duyệt qua các trang để chạy OMR
        for page_num, img in pages_to_process:
            sheet_id = str(uuid.uuid4())
            try:
                sheet_data, err, debug_img = processor.process_sheet(img)
                
                if err:
                    # Ghi nhận lỗi quét nhưng vẫn cho vào danh sách để người dùng có thể tải lại hoặc sửa bằng tay
                    sheet_result = {
                        "sheet_id": sheet_id,
                        "filename": filename,
                        "page": page_num,
                        "cccd": "",
                        "de_thi": "",
                        "student_info": None,
                        "answers": {str(q): "TRONG" for q in range(1, 31)},
                        "correct_count": 0,
                        "score": 0,
                        "has_error": True,
                        "error_message": err,
                        "debug_img": None
                    }
                else:
                    cccd = sheet_data["cccd"]
                    dethi = sheet_data.get("de_thi", "")
                    answers = sheet_data["answers"]
                    
                    # Khớp tên từ danh sách
                    student_info = CANDIDATES_DB.get(cccd, None)
                    correct_count, score = calculate_score(answers)
                    
                    # Xác định xem bài quét có lỗi không (bao gồm CCCD không hợp lệ hoặc thiếu đáp án)
                    has_error = (
                        student_info is None or 
                        "?" in cccd or 
                        "?" in dethi or
                        any(v in ["TRONG", "TRUNG"] for v in answers.values())
                    )
                    
                    error_msg = ""
                    if "?" in cccd:
                        error_msg = "Lỗi tô CCCD. "
                    elif student_info is None:
                        error_msg = "Không tìm thấy CCCD trong danh sách sinh viên. "
                        
                    if "?" in dethi:
                        error_msg += "Lỗi tô Đề thi số. "
                        
                    if any(v == "TRONG" for v in answers.values()):
                        error_msg += "Có câu chưa tô. "
                    if any(v == "TRUNG" for v in answers.values()):
                        error_msg += "Có câu tô trùng đáp án. "
                        
                    sheet_result = {
                        "sheet_id": sheet_id,
                        "filename": filename,
                        "page": page_num,
                        "cccd": cccd,
                        "de_thi": dethi,
                        "student_info": student_info,
                        "answers": answers,
                        "correct_count": correct_count,
                        "score": score,
                        "has_error": has_error,
                        "error_message": error_msg.strip(),
                        "debug_img": debug_img
                    }
                
                # Lưu vào bộ nhớ tạm thời của Server
                PROCESSED_SHEETS[sheet_id] = sheet_result
                
                # Để giảm dung lượng JSON trả về, ta clone và loại bỏ ảnh debug nếu không lỗi hoặc tùy nhu cầu
                # Tuy nhiên để hiển thị sửa nhanh thì vẫn giữ lại debug_img trong PROCESSED_SHEETS
                client_item = {k: v for k, v in sheet_result.items() if k != "debug_img"}
                results.append(client_item)
                
            except Exception as e:
                results.append({
                    "filename": filename,
                    "page": page_num,
                    "error": f"Lỗi hệ thống khi chấm: {str(e)}"
                })
                
        # Xóa file đã lưu sau khi xử lý xong
        try:
            os.remove(file_path)
        except Exception:
            pass
            
    return jsonify({"status": "success", "results": results})

@app.route("/api/sheet/<sheet_id>", methods=["GET"])
def get_sheet_detail(sheet_id):
    sheet = PROCESSED_SHEETS.get(sheet_id)
    if not sheet:
        return jsonify({"status": "error", "message": "Không tìm thấy phiếu này trong bộ nhớ."}), 404
        
    # Clone sheet kết quả để vẽ thêm đè lên ảnh hiển thị
    sheet_clone = json.loads(json.dumps(sheet))
    sheet_clone["answer_key"] = ANSWER_KEY
    
    if sheet_clone.get("debug_img") and ANSWER_KEY:
        try:
            import base64
            # Giải mã ảnh debug_img thành OpenCV Image
            img_data = base64.b64decode(sheet_clone["debug_img"])
            nparr = np.frombuffer(img_data, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            
            if img is not None:
                # Vẽ vòng tròn đỏ rực to hơn (bán kính 12, độ dày 2) bao quanh đáp án đúng
                options_map = {"A": 0, "B": 1, "C": 2, "D": 3}
                for q_str, correct_ans in ANSWER_KEY.items():
                    if q_str in processor.coords["answers"]:
                        opt_idx = options_map.get(correct_ans)
                        if opt_idx is not None:
                            cx, cy = processor.coords["answers"][q_str][opt_idx]
                            # BGR: màu đỏ (0, 0, 255)
                            cv2.circle(img, (cx, cy), 12, (0, 0, 255), 2)
                            
                # Lưu lại ảnh đã vẽ đè vòng tròn đỏ
                _, buffer = cv2.imencode('.jpg', img)
                sheet_clone["debug_img"] = base64.b64encode(buffer).decode('utf-8')
        except Exception as e:
            print(f"[!] Lỗi khi vẽ đè đáp án đúng lên debug_img: {e}")
            
    return jsonify(sheet_clone)

@app.route("/api/sheet/<sheet_id>", methods=["POST"])
def update_sheet_detail(sheet_id):
    """Cập nhật lại CCCD hoặc câu trả lời sau khi sửa tay trên giao diện"""
    sheet = PROCESSED_SHEETS.get(sheet_id)
    if not sheet:
        return jsonify({"status": "error", "message": "Không tìm thấy phiếu."}), 404
        
    data = request.json
    cccd = data.get("cccd", "").strip()
    dethi = data.get("de_thi", "").strip()
    answers = data.get("answers", {})
    
    # Nhận thông tin thí sinh mới nếu có nhập tay
    name = data.get("name", "").strip()
    dob = data.get("dob", "").strip()
    address = data.get("address", "").strip()
    
    if name and cccd and len(cccd) == 12:
        # Thêm mới thí sinh vào cơ sở dữ liệu tạm thời
        CANDIDATES_DB[cccd] = {
            "name": name,
            "dob": dob,
            "date_issued": "",
            "address": address
        }
        # Lưu đè bổ sung vào tệp candidates.xlsx gốc
        excel_path = os.path.join(BASE_DIR, "data", "candidates.xlsx")
        try:
            import pandas as pd
            # Đọc các sheets hiện tại
            xls = pd.ExcelFile(excel_path)
            sheets = xls.sheet_names
            dict_dfs = {s: xls.parse(s) for s in sheets}
            
            # Thêm dòng mới vào sheet đầu tiên
            first_sheet = sheets[0]
            df_first = dict_dfs[first_sheet]
            
            new_row = {
                "Tên": name,
                "Năm sinh": dob,
                "CCCD": cccd,
                "Ngày cấp": "",
                "Nơi thường trú": address
            }
            # Append dòng mới
            df_first = pd.concat([df_first, pd.DataFrame([new_row])], ignore_index=True)
            dict_dfs[first_sheet] = df_first
            
            # Ghi lại toàn bộ file Excel
            with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
                for s_name, df_sheet in dict_dfs.items():
                    df_sheet.to_excel(writer, sheet_name=s_name, index=False)
            print(f"[+] Tự động thêm thí sinh mới vào Excel: {cccd} - {name}")
        except Exception as e:
            print(f"[!] Không thể ghi thêm thí sinh vào Excel: {e}")

    # Cập nhật thông tin mới
    sheet["cccd"] = cccd
    sheet["de_thi"] = dethi
    sheet["answers"].update(answers)
    
    # Tính toán lại điểm số và kiểm tra danh sách thí sinh
    student_info = CANDIDATES_DB.get(cccd, None)
    correct_count, score = calculate_score(sheet["answers"])
    
    sheet["student_info"] = student_info
    sheet["correct_count"] = correct_count
    sheet["score"] = score
    
    # Kiểm tra xem còn lỗi hay không
    has_error = (
        student_info is None or 
        "?" in cccd or 
        "?" in dethi or
        any(v in ["TRONG", "TRUNG"] for v in sheet["answers"].values())
    )
    
    error_msg = ""
    if "?" in cccd or not cccd:
        error_msg = "Lỗi/Thiếu CCCD. "
    elif student_info is None:
        error_msg = "Không tìm thấy CCCD trong danh sách sinh viên. "
        
    if "?" in dethi or not dethi:
        error_msg += "Lỗi/Thiếu Đề thi số. "
        
    if any(v == "TRONG" for v in sheet["answers"].values()):
        error_msg += "Có câu chưa tô. "
    if any(v == "TRUNG" for v in sheet["answers"].values()):
        error_msg += "Có câu tô trùng đáp án. "
        
    sheet["has_error"] = has_error
    sheet["error_message"] = error_msg.strip()
    
    # Trả về kết quả sau khi cập nhật (loại bỏ debug_img để tối ưu traffic)
    client_item = {k: v for k, v in sheet.items() if k != "debug_img"}
    return jsonify({"status": "success", "sheet": client_item})

@app.route("/api/export-excel", methods=["GET"])
def export_excel():
    """
    Đọc tệp danh sách thí sinh Excel gốc tại data/candidates.xlsx,
    ghi thêm điểm số cho từng thí sinh tương ứng trong từng sheet,
    và bổ sung thêm sheet 'Thí sinh ngoài danh sách' cho các bài quét tự do.
    """
    excel_path = os.path.join(BASE_DIR, "data", "candidates.xlsx")
    
    # Gom kết quả chấm theo CCCD
    results_map = {}
    for sheet in PROCESSED_SHEETS.values():
        cccd = sheet.get("cccd", "")
        if cccd and not sheet.get("has_error"):
            ans_str = "|".join([f"{q}:{a}" for q, a in sorted(sheet["answers"].items(), key=lambda x: int(x[0]))])
            results_map[cccd] = {
                "de_thi": sheet.get("de_thi", ""),
                "score": sheet["score"],
                "correct_count": sheet["correct_count"],
                "answers": ans_str
            }

    try:
        import pandas as pd
        import io
        
        output = io.BytesIO()
        
        if os.path.exists(excel_path):
            xls = pd.ExcelFile(excel_path)
            sheets = xls.sheet_names
            dict_dfs = {}
            written_cccds = set()
            
            for s_name in sheets:
                df = xls.parse(s_name)
                # Làm sạch headers
                df.columns = [str(c).strip() for c in df.columns]
                
                if "CCCD" in df.columns:
                    col_dethi = []
                    col_correct = []
                    col_score = []
                    col_answers = []
                    
                    for _, row in df.iterrows():
                        cccd_raw = str(row["CCCD"]).strip()
                        cccd = cccd_raw.split(".")[0] if cccd_raw and cccd_raw != "nan" else ""
                        
                        if cccd in results_map:
                            res = results_map[cccd]
                            col_dethi.append(res["de_thi"])
                            col_correct.append(f"{res['correct_count']}/30")
                            col_score.append(res["score"])
                            col_answers.append(res["answers"])
                            written_cccds.add(cccd)
                        else:
                            col_dethi.append("")
                            col_correct.append("Chưa làm bài")
                            col_score.append(0.0)
                            col_answers.append("")
                            
                    df["Đề thi số"] = col_dethi
                    df["Số câu đúng / 30"] = col_correct
                    df["Điểm số"] = col_score
                    df["Đáp án chi tiết"] = col_answers
                    
                dict_dfs[s_name] = df
                
            # Thêm sheet phụ nếu có thí sinh ngoài danh sách
            free_candidates = []
            for cccd, res in results_map.items():
                if cccd not in written_cccds:
                    free_candidates.append({
                        "Tên": f"Thí sinh tự do (CCCD: {cccd})",
                        "CCCD": cccd,
                        "Đề thi số": res["de_thi"],
                        "Số câu đúng / 30": f"{res['correct_count']}/30",
                        "Điểm số": res["score"],
                        "Đáp án chi tiết": res["answers"]
                    })
                    
            if free_candidates:
                df_free = pd.DataFrame(free_candidates)
                dict_dfs["Thí sinh ngoài danh sách"] = df_free
                
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                for s_name, df_sheet in dict_dfs.items():
                    df_sheet.to_excel(writer, sheet_name=s_name, index=False)
        else:
            # Fallback tạo bảng Excel tổng hợp nếu file candidates.xlsx không tồn tại
            all_results = []
            for cccd, res in results_map.items():
                student_name = "Thí sinh tự do"
                student_info = CANDIDATES_DB.get(cccd)
                if student_info and student_info.get("name"):
                    student_name = student_info["name"]
                    
                all_results.append({
                    "Tên": student_name,
                    "CCCD": cccd,
                    "Đề thi số": res["de_thi"],
                    "Số câu đúng / 30": f"{res['correct_count']}/30",
                    "Điểm số": res["score"],
                    "Đáp án chi tiết": res["answers"]
                })
            df_all = pd.DataFrame(all_results)
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df_all.to_excel(writer, sheet_name="Kết quả chấm thi", index=False)
                
        output.seek(0)
        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="ket_qua_cham_thi.xlsx"
        )
    except Exception as e:
        return jsonify({"status": "error", "message": f"Lỗi xuất file Excel: {str(e)}"}), 500

@app.route("/api/export-docx", methods=["GET"])
def export_docx():
    """
    Tạo file Word (.docx) chứa điểm số căn chỉnh đúng vị trí ô điểm số trên giấy thi A4.
    Mỗi trang tương ứng với một phiếu đã chấm.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    # Lấy tất cả các phiếu đã chấm (không loại bỏ lỗi để giữ nguyên thứ tự in giấy thi)
    sheets_list = list(PROCESSED_SHEETS.values())
            
    if not sheets_list:
        return jsonify({"status": "error", "message": "Chưa có bài thi nào được chấm để xuất điểm."}), 400
        
    try:
        doc = Document()
        
        # Cấu hình lề trang đầu tiên
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.5) # Căn chỉnh thẳng vào ô chữ nhật y=571px
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Helper định dạng điểm số tiếng Việt
        def format_score(score):
            s = f"{score:.2f}"
            if s.endswith(".00"):
                return s[:-3] + ".0"
            return s

        for idx, sheet in enumerate(sheets_list):
            if idx > 0:
                doc.add_page_break()
                
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            score_text = format_score(sheet["score"])
            run = p.add_run(score_text)
            run.font.name = "Arial"
            run.font.size = Pt(48)
            run.font.bold = True
            run.font.color.rgb = RGBColor(220, 38, 38) # Màu đỏ rực rỡ để in rõ ràng
            
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="in_diem_thi.docx"
        )
    except Exception as e:
        return jsonify({"status": "error", "message": f"Lỗi xuất file Word: {str(e)}"}), 500

@app.route("/api/upload-template", methods=["POST"])
def upload_template():
    if 'template_file' not in request.files:
        return jsonify({"status": "error", "message": "Không tìm thấy file tải lên."}), 400
        
    file = request.files['template_file']
    if not file.filename:
        return jsonify({"status": "error", "message": "File không hợp lệ."}), 400
        
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], "template_" + filename)
    file.save(file_path)
    
    _, ext = os.path.splitext(filename.lower())
    
    try:
        # Chuyển đổi sang ảnh nếu là PDF
        if ext == '.pdf':
            doc = fitz.open(file_path)
            if len(doc) == 0:
                return jsonify({"status": "error", "message": "File PDF rỗng."}), 400
            page = doc[0]
            pix = page.get_pixmap(dpi=300)
            img_data = pix.tobytes("png")
            nparr = np.frombuffer(img_data, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        else:
            img = cv2.imread(file_path)
            
        if img is None:
            return jsonify({"status": "error", "message": "Không đọc được dữ liệu ảnh."}), 400
            
        # Tìm các góc định vị
        anchors = processor.find_anchors(img)
        if len(anchors) < 4:
            return jsonify({
                "status": "error", 
                "message": f"Tìm thấy {len(anchors)}/4 điểm định vị. Phiếu của bạn phải chứa đủ 4 hình vuông màu đen đặc ở 4 góc để hệ thống căn thẳng ảnh."
            }), 400
            
        sorted_pts = processor.sort_anchors(anchors)
        TL, TR, BL, BR = sorted_pts
        
        src_pts = np.float32([TL, TR, BL, BR])
        dst_pts = np.float32([[50, 50], [950, 50], [50, 1364], [950, 1364]])
        
        M = cv2.getPerspectiveTransform(src_pts, dst_pts)
        warped = cv2.warpPerspective(img, M, (1000, 1414))
        
        # Encode sang base64 để gửi lên giao diện
        _, buffer = cv2.imencode('.jpg', warped)
        import base64
        warped_base64 = base64.b64encode(buffer).decode('utf-8')
        
        # Xóa file tạm
        try:
            os.remove(file_path)
        except Exception:
            pass
            
        return jsonify({
            "status": "success", 
            "warped_image": warped_base64
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": f"Lỗi hệ thống khi xử lý mẫu phiếu: {str(e)}"}), 500

@app.route("/api/save-coordinates", methods=["POST"])
def save_coordinates():
    global processor
    data = request.json
    
    # Kiểm tra tính hợp lệ sơ bộ
    if not data or "cccd" not in data or "de_thi" not in data or "answers" not in data:
        return jsonify({"status": "error", "message": "Dữ liệu cấu hình tọa độ không hợp lệ."}), 400
        
    coord_json_path = os.path.join(BASE_DIR, "data", "bubble_coordinates.json")
    
    try:
        with open(coord_json_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)
            
        # Re-initialize processor để ăn tọa độ mới ngay lập tức
        processor = OMRProcessor(coord_json_path)
        print("[+] Coordinates updated and re-loaded successfully.")
        
        # Tự động gọi script tái tạo lại 5 bài mẫu khớp với tọa độ vừa lưu
        try:
            import subprocess
            import sys
            python_bin = sys.executable
            script_path = os.path.join(BASE_DIR, "simulate_filled_sheet.py")
            subprocess.run([python_bin, script_path], check=True)
            print("[+] Automatically regenerated 5 sample sheets matching new coordinates!")
        except Exception as e:
            print(f"[!] Error auto-regenerating sample sheets: {e}")
            
        return jsonify({"status": "success", "message": "Đã lưu tọa độ mới và tự động tái tạo 5 phiếu làm bài mẫu thành công!"})
    except Exception as e:
        return jsonify({"status": "error", "message": f"Không thể lưu cấu hình tọa độ: {str(e)}"}), 500

@app.route("/api/current-calibration", methods=["GET"])
def current_calibration():
    pdf_path = os.path.join(BASE_DIR, "phieu_trac_nghiem.pdf")
    png_template_path = os.path.join(BASE_DIR, "mau_thiet_ke_raw.png")
    coord_json_path = os.path.join(BASE_DIR, "data", "bubble_coordinates.json")
    
    if not os.path.exists(coord_json_path):
        return jsonify({"status": "error", "message": "Không tìm thấy tệp tọa độ mẫu."}), 404
        
    try:
        # 1. Đọc ảnh mẫu thiết kế (ưu tiên tệp PNG Photoshop của người dùng nếu có)
        if os.path.exists(png_template_path):
            img = cv2.imread(png_template_path)
            print("[*] Đang tải mẫu thiết kế Photoshop PNG:", png_template_path)
        elif os.path.exists(pdf_path):
            doc = fitz.open(pdf_path)
            page = doc[0]
            pix = page.get_pixmap(dpi=300, alpha=False)
            img_data = pix.tobytes("png")
            nparr = np.frombuffer(img_data, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        else:
            return jsonify({"status": "error", "message": "Không tìm thấy tệp mẫu thiết kế PDF hay PNG."}), 404
            
        if img is None:
            return jsonify({"status": "error", "message": "Không thể giải mã hình ảnh mẫu."}), 400
            
        anchors = processor.find_anchors(img)
        if len(anchors) < 4:
            return jsonify({
                "status": "error", 
                "message": f"Tìm thấy {len(anchors)}/4 góc đen định vị trên ảnh Photoshop. Vui lòng kiểm tra xem bạn có vô tình che mất hoặc di chuyển 4 hình vuông đen ở 4 góc không."
            }), 400
            
        sorted_pts = processor.sort_anchors(anchors)
        TL, TR, BL, BR = sorted_pts
        src_pts = np.float32([TL, TR, BL, BR])
        dst_pts = np.float32([[50, 50], [950, 50], [50, 1364], [950, 1364]])
        
        M = cv2.getPerspectiveTransform(src_pts, dst_pts)
        warped = cv2.warpPerspective(img, M, (1000, 1414))
        
        _, buffer = cv2.imencode('.jpg', warped)
        import base64
        warped_base64 = base64.b64encode(buffer).decode('utf-8')
        
        # 2. Đọc tọa độ hiện tại và trích xuất 10 điểm mốc
        with open(coord_json_path, "r", encoding="utf-8") as f:
            coords = json.load(f)
            
        points = {
            "cccd_tl": coords["cccd"][0][0],
            "cccd_br": coords["cccd"][11][9],
            "dethi_tl": coords["de_thi"][0][0],
            "dethi_br": coords["de_thi"][3][9],
            "q1_10_tl": coords["answers"]["1"][0],
            "q1_10_br": coords["answers"]["10"][3],
            "q11_20_tl": coords["answers"]["11"][0],
            "q11_20_br": coords["answers"]["20"][3],
            "q21_30_tl": coords["answers"]["21"][0],
            "q21_30_br": coords["answers"]["30"][3]
        }
        
        return jsonify({
            "status": "success",
            "image": warped_base64,
            "points": points
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

if __name__ == "__main__":
    # Load lại dữ liệu đề phòng file thay đổi
    load_candidates()
    load_answer_key()
    app.run(host="127.0.0.1", port=5000, debug=True)
